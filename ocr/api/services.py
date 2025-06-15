import os
import base64
import json
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Pt
from typing import List
from django.core.files.storage import FileSystemStorage
from django.core.files.uploadedfile import UploadedFile
import firebase_admin
from firebase_admin import credentials, firestore

from application.services import paddle_model, easy_model

from application.model.paddleocr import PaddleOCR
from application.model.easyocr import EasyOCR
from application.utils import validate_image_brightness

paddle_model = PaddleOCR()
easy_model = EasyOCR()

DEBUG_MODE = False
POLISH_MODE = False


def get_files():
    """
    Retrieve a list of all files from the configured upload directory.
    :return: List of filenames in the upload directory, or ["empty"] if directory not found.
    """
    directory = os.environ.get("UPLOADED_FILES")
    try:
        return os.listdir(directory)
    except FileNotFoundError:
        return ["empty"]


def prepare_file_hierarchy(file):
    """
    Save an uploaded file to the designated upload directory and return its full path.
    :param file: The uploaded file object to be saved.
    :return: The full file system path where the file was saved.
    """
    upload_dir = os.path.abspath(os.environ["UPLOADED_FILES"])
    if DEBUG_MODE:
        print(f"Upload directory: {upload_dir}")

    # Save the uploaded file first
    storage = FileSystemStorage(location=upload_dir)
    file_path = storage.save(file.name, file)
    full_path = storage.path(file_path)
    if DEBUG_MODE:
        print(f"Saved file to: {full_path}")

    return full_path


def convert_result_to_json(
    uploaded_file: UploadedFile,
    line_width: int,
    font_size: int,
    uploaded_file_path: str,
    ocr_result: dict,
):
    """
    Convert OCR results to JSON format for frontend consumption with base64 encoded image data.
    :param uploaded_file: Django UploadedFile object from request.FILES["file"].
    :param line_width: The character length limit for lines in potential .txt output files.
    :param font_size: The font size for text in potential .docx output files.
    :param uploaded_file_path: Absolute file system path where the uploaded file was saved.
    :param ocr_result: Dictionary containing OCR results in standardized ModelBase format.
    :return: JSON-formatted string containing file name, base64 image data, format, confidence scores, and OCR content.
    """
    name, file_format = uploaded_file.name.split(".")

    img = Image.open(uploaded_file_path)
    im_file_in_mem = BytesIO()

    if file_format == "jpg" or file_format == "jpeg":
        img.save(im_file_in_mem, format="JPEG")
    elif file_format == "png":
        img.save(im_file_in_mem, format="PNG")

    im_bytes = im_file_in_mem.getvalue()
    im_b64 = base64.b64encode(im_bytes).decode("utf-8")

    confidence_list = ocr_result["confidence_scores"]
    content_list = ocr_result["text_predictions"]

    json_str = json.dumps(
        {
            "name": name,
            "image": im_b64,
            "format": file_format,
            "paragraphWidth": line_width,
            "fontSize": font_size,
            "confidence": confidence_list,
            "content": content_list,
        }
    )

    return json_str


def handle_uploaded_file(file, file_format, line_width, font_size, user_uid=None):
    """
    Process an uploaded file by saving it, validating brightness, and performing OCR based on language mode.
    :param file: The uploaded file object to be processed.
    :param file_format: The desired output format for the processed text.
    :param line_width: The character length limit for lines in potential .txt output files.
    :param font_size: The font size for text in potential .docx output files.
    :param user_uid: Optional user identifier for associating the file with a specific user.
    """
    full_path = prepare_file_hierarchy(file)

    if validate_image_brightness(full_path):
        if POLISH_MODE:
            result = easy_model.perform_ocr(input_path=full_path)
            if DEBUG_MODE:
                print("EasyOCR results:")
                print(" ".join([word for word in result["text_predictions"]]))

        if not POLISH_MODE:
            result = paddle_model.perform_ocr(input_path=full_path)
            if DEBUG_MODE:
                print("PaddleOCR results:")
                print(" ".join([word for word in result["text_predictions"]]))

        # Initialize Firebase if not already initialized
        if not firebase_admin._apps:
            cred_path = os.environ["FIREBASE_KEY"]
            if not os.path.exists(cred_path):
                raise Exception("Firebase credentials not found")
            cred = credentials.Certificate(cred_path)
            firebase_admin.initialize_app(cred)

        # Get Firestore client
        db = firestore.client()

        # Read and encode the image
        with open(full_path, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode("utf-8")

        # Add image to Firestore with userUid and OCR results
        image_ref = db.collection("images").document()
        image_ref.set(
            {
                "image_data": encoded_string,
                "filename": file.name,
                "timestamp": firestore.SERVER_TIMESTAMP,
                "userUid": user_uid,  # This will be null for unauthenticated users
                "ocr_results": {
                    "text_predictions": result["text_predictions"],
                    "confidence_scores": result["confidence_scores"],
                },
                "format": file_format,
                "paragraphWidth": line_width,
                "fontSize": font_size,
            }
        )

        json_ocr_result = convert_result_to_json(
            file, line_width, font_size, full_path, result
        )

        return json_ocr_result


def output_processed_as_txt(word_list: List[str], output_path: str, line_width=80):
    """
    Output a list of words (strings) extracted from an image into a .txt file at output_path.
    :param word_list: List of strings which are extracted words.
    :param output_path: A string representing the output_path for the .txt file to be generated at.
    :param line_width: The limit in characters after which a full new word cannot be written (but a part of the word may).
    """
    chars_in_current_line = 0
    max_line_width = line_width
    with open(output_path, "w") as f:
        for word in word_list:
            chars_in_current_line += len(word)
            # If limit exceeded after writing this word, write
            if chars_in_current_line > max_line_width:
                # But write with a newline and reset next line to the beginning
                f.writelines(f"{word}\n")
                chars_in_current_line = 0
            else:
                # Otherwise simply write
                f.writelines(f"{word} ")


def output_processed_as_docx(word_list: List[str], output_path: str, font_size=11):
    """
    Output a list of words (strings) extracted from an image into a .docx file at output_path.
    :param word_list: List of strings which are extracted words.
    :param output_path: A string representing the output_path for the .txt file to be generated at.
    :param font_size: The size of the font of the text which will take up the entire width at all times.
    """
    document = Document()

    # Add_run is the most general "generate text" method. It gives access to a Font obj through .font
    font = document.add_paragraph().add_run(" ".join(word_list)).font

    # Through the Font obj we can modify the appearance of the text
    font.name = "Arial"
    font.size = Pt(font_size)

    # The appearance will be saved
    document.save(output_path)
