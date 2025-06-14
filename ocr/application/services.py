import os
from docx import Document
from docx.shared import Pt
from typing import List

from pathlib import Path
import base64
from io import BytesIO

import firebase_admin
from firebase_admin import credentials, firestore
from google.cloud.firestore_v1 import FieldFilter

from PIL import Image
from typing import List
from django.core.files.storage import FileSystemStorage

from application.model.paddleocr import PaddleOCR
from application.model.easyocr import EasyOCR
from application.utils import validate_image_brightness

paddle_model = PaddleOCR()
easy_model = EasyOCR()

firestore_db = None


def prepare_file_hierarchy(file):
    """Takes uploaded file and returns directory where it is saved and its detected content"""
    upload_dir = os.path.abspath(os.environ['UPLOADED_FILES'])
    print(f"Upload directory: {upload_dir}")

    # Save the uploaded file first
    storage = FileSystemStorage(location=upload_dir)
    file_path = storage.save(file.name, file)
    full_path = storage.path(file_path)
    print(f"Saved file to: {full_path}")

    # Possibly legacy?
    # output_dir = os.path.join(upload_dir, 'processed_text')
    # os.makedirs(output_dir, exist_ok=True)

    return full_path

def handle_uploaded_file(file):
    full_path = prepare_file_hierarchy(file)

    try:
        if not validate_image_brightness(full_path):
            return {'status': 'error', 'message': 'Image too dark.'}

        paddle_result = paddle_model.perform_ocr(input_path=full_path)
        print("PaddleOCR results:")
        print(" ".join([result for result in paddle_result["text_predictions"]]))

        easy_result = easy_model.perform_ocr(input_path=full_path)
        print("EasyOCR results:")
        print(" ".join([result for result in easy_result["text_predictions"]]))

        combined_text = paddle_result["text_predictions"] + easy_result["text_predictions"]
        if not combined_text:
            return {'status': 'error', 'message': 'No text could be extracted from the image.'}

        return {
            'status': 'success',
            'text': " ".join(combined_text)
        }

    except Exception as e:
        return {
            'status': 'error',
            'message': str(e)
        }

def setup_firestore_db():
    """
    Initialize the global firestore_db variable to hold a google.cloud.firestore_v1 Client object.
    """
    # There may be a more elegant solution, but for now I am using a global variable
    global firestore_db

    # The environment variable is pulled from rootdir/ocr/.env
    credentials_obj = credentials.Certificate(os.environ["FIREBASE_KEY"])
    firebase_admin.initialize_app(credentials_obj)
    firestore_db = firestore.client()

def retrieve_pictures_using_uid(desired_uid: str) -> List[Image.Image]:
    """
    Retrieve a list of a user's PIL.Image objects from the Firestore database based on his UID.
    Depends on the global (services.py) variable firebase_db being initialized. This variable is
    initialized with setup_firestore_db (services.py) called within ApplicationConfig (apps.py).

    :param desired_uid: The UID of the user for which pictures are to be retrieved from the db.
    :return: A list of Pillow image objects generated from parsing image data stored in the db.
    """

    # A list of DocumentSnapshot objects meeting the userUid == desired_uid criterion
    user_images_snapshot_list = (
        firestore_db.collection("images")
        .where(filter=FieldFilter("userUid", "==", desired_uid))
        .stream()
    )

    # A list of user images encoded in b64 string format
    user_images_b64_list = [user_image_snapshot.get("image_data") for user_image_snapshot in user_images_snapshot_list]

    # A list of PIL.Image objects generated from decoding the b64 string format
    user_images_list = [Image.open(BytesIO(base64.b64decode(user_image_b64))) for user_image_b64 in user_images_b64_list]

    return user_images_list

def output_processed_as_txt(word_list: List[str], output_path: str, line_width=80):
    """
    Output a list of words (strings) extracted from an image into a .txt file at output_path.
    :param word_list: List of strings which are extracted words.
    :param output_path: A string representing the output_path for the .txt file to be generated at.
    :param line_width: The limit in characters after which a full new word cannot be written (but a part of the word may).
    """
    chars_in_current_line = 0
    max_line_width = line_width
    with open(output_path, "w") as f:
        for word in word_list:
            chars_in_current_line += len(word)
            # If limit exceeded after writing this word, write
            if chars_in_current_line > max_line_width:
                # But write with a newline and reset next line to the beginning
                f.writelines(f"{word}\n")
                chars_in_current_line = 0
            else:
                # Otherwise simply write
                f.writelines(f"{word} ")

def output_processed_as_docx(word_list: List[str], output_path: str, font_size=11):
    """
    Output a list of words (strings) extracted from an image into a .docx file at output_path.
    :param word_list: List of strings which are extracted words.
    :param output_path: A string representing the output_path for the .txt file to be generated at.
    :param font_size: The size of the font of the text which will take up the entire width at all times.
    """
    document = Document()

    # Add_run is the most general "generate text" method. It gives access to a Font obj through .font
    font = document.add_paragraph().add_run(" ".join(word_list)).font

    # Through the Font obj we can modify the appearance of the text
    font.name = "Arial"
    font.size = Pt(font_size)

    # The appearance will be saved
    document.save(output_path)

def get_db():
    if not firebase_admin._apps:
        cred_path = os.environ['FIREBASE_KEY']
        if not os.path.exists(cred_path):
            raise Exception('Firebase credentials not found')
        cred = credentials.Certificate(cred_path)
        firebase_admin.initialize_app(cred)

    return firestore.client()

def set_data_limit(data_limit):
    db = get_db()
    db.collection("global_settings").document("limits").set({
        'dataLimit': data_limit,
    }, merge=True)

def set_file_limit(file_limit):
    db = get_db()
    db.collection("global_settings").document("limits").set({
        'fileLimit': file_limit,
    }, merge=True)

def get_limits():
    db = get_db()
    doc = db.collection("global_settings").document("limits").get()
    return doc.to_dict() if doc.exists else {}